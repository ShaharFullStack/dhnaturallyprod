import os
import re
import uuid
from docx import Document

PRODUCTS_DIR = os.path.join(os.path.dirname(__file__), '..', 'Database', 'products')
OUTPUT_SQL = os.path.join(os.path.dirname(__file__), '..', 'Database', 'workbench_import.sql')

# Basic heuristics to extract a title and the first meaningful paragraph

def extract_text_from_docx(path):
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    if not paragraphs:
        return None, None
    # Use filename (without extension) as fallback title
    title = paragraphs[0] if len(paragraphs) == 1 else paragraphs[0]
    # find first paragraph longer than 40 chars as description, else second paragraph
    desc = None
    for p in paragraphs[1:6]:
        if len(p) > 40:
            desc = p
            break
    if not desc and len(paragraphs) > 1:
        desc = paragraphs[1]
    return title, desc


def contains_hebrew(s):
    if not s:
        return False
    return bool(re.search(r"[\u0590-\u05FF]", s))


def sanitize_sql(s):
    if s is None:
        return ''
    return s.replace("'", "''")


def main():
    entries = []
    for fname in sorted(os.listdir(PRODUCTS_DIR)):
        if not fname.lower().endswith('.docx'):
            continue
        path = os.path.join(PRODUCTS_DIR, fname)
        title, desc = extract_text_from_docx(path)
        fallback_name = os.path.splitext(fname)[0]
        name_en = title if title else fallback_name
        desc_en = desc if desc else ''
        # If doc contains Hebrew, try to separate; simple heuristic: if contains Hebrew, use whole text as Hebrew desc
        full_text = '\n'.join([p for p in Document(path).paragraphs if p.text and p.text.strip()])
        if contains_hebrew(full_text):
            name_he = name_en  # no good heuristic for hebrew title; keep english name
            desc_he = full_text
        else:
            name_he = name_en
            desc_he = 'תיאור בעברית לא זמין; יש לתרגם'
        entries.append({
            'name_en': sanitize_sql(name_en),
            'name_he': sanitize_sql(name_he),
            'desc_en': sanitize_sql(desc_en),
            'desc_he': sanitize_sql(desc_he),
        })

    # Build SQL file header
    header = """-- DHnaturally Workbench Import Script
-- Creates database, products table and inserts sample products
-- Ready to paste into MySQL Workbench or run from the mysql client

-- 1) Create database with UTF8MB4 for Hebrew support
CREATE DATABASE IF NOT EXISTS dhnaturally_db
CHARACTER SET utf8mb4 
COLLATE utf8mb4_unicode_ci;

USE dhnaturally_db;

-- 2) Create a minimal products table compatible with provided seed data
CREATE TABLE IF NOT EXISTS products (
    id VARCHAR(36) PRIMARY KEY,
    name VARCHAR(200) NOT NULL,
    name_he VARCHAR(200) NOT NULL,
    description TEXT NOT NULL,
    description_he TEXT NOT NULL,
    price DECIMAL(10,2) NOT NULL,
    imageName VARCHAR(255),
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
    INDEX idx_name (name),
    INDEX idx_price (price),
    INDEX idx_created_at (created_at)
);

INSERT INTO products (id, name, name_he, description, description_he, price, imageName) VALUES
"""

    lines = [header]
    for i, e in enumerate(entries):
        comma = ',' if i < len(entries)-1 else ';'
        line = "(UUID(), '{name_en}', '{name_he}', '{desc_en}', '{desc_he}', 169.90, 'product.png')".format(**e) + comma
        lines.append(line)

    lines.append('\n-- 4) Ensure foreign key checks are enabled (if you run other scripts that require them)')
    lines.append('SET FOREIGN_KEY_CHECKS = 1;')
    lines.append('\n-- End of script')

    with open(OUTPUT_SQL, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))

    print(f'Wrote {len(entries)} products to {OUTPUT_SQL}')


if __name__ == '__main__':
    main()
